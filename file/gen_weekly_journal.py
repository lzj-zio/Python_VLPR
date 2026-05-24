# -*- coding: utf-8 -*-
"""从 journal_data.txt 读取内容，生成毕业设计过程周记 DOCX"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os, re

# 读取数据文件
data_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "journal_data.txt")
with open(data_path, "r", encoding="utf-8") as f:
    raw = f.read().strip()

# 按 === 分割每篇周记
blocks = [b.strip() for b in raw.split("\n===\n") if b.strip()]
print(f"共读取 {len(blocks)} 篇周记")

# 解析每篇周记
journals = []
for block in blocks:
    lines = block.split("\n")
    # 第一行格式：第X周|日期|标题
    header = lines[0]
    parts = header.split("|")
    week = parts[0].strip()
    date = parts[1].strip()
    title = parts[2].strip()
    content = "\n".join(lines[1:]).strip()
    journals.append({
        "week": week,
        "date": date,
        "title": title,
        "content": content,
    })

# 生成 DOCX
doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.17)
section.right_margin = Cm(3.17)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

for j in journals:
    # 周次标题
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(f"{j['week']}周记")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 日期+标题行
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_info.add_run(f"日期：{j['date']}    {j['title']}")
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 正文（按空行分段）
    paragraphs = j['content'].split('\n\n')
    for pi, para_text in enumerate(paragraphs):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.74)
        run = p.add_run(para_text.strip())
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 字数统计
    char_count = len(j['content'].replace('\n', '').replace(' ', ''))
    print(f"  {j['week']} [{j['title']}] - {char_count}字")

    # 分页符（最后一篇不加）
    if j != journals[-1]:
        doc.add_page_break()

# 保存
output_path = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    "毕业设计过程周记_李志杰_2022b15080.docx"
)
doc.save(output_path)
print(f"\n文档已保存: {output_path}")
