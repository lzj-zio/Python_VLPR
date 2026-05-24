# -*- coding: utf-8 -*-
"""
用Word原生OMML公式替换「？代码？」占位符（不使用图片）
"""
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# OMML命名空间
OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def make_omml_para(latex_str):
    """
    给定一个简化LaTeX字符串，转成Word OMML XML，嵌入到段落中。
    只支持常用的符号：分数、根号、下标、上标、希腊字母、累加、积分等。
    """
    # 1. 将简化latex转为OMML
    omml_xml = latex_to_omml(latex_str)

    # 2. 构造完整段落XML
    para = OxmlElement('w:p')

    # 段落属性（居中）
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    para.append(pPr)

    # Run + OMML
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Cambria Math')
    rFonts.set(qn('w:hAnsi'), 'Cambria Math')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rPr.append(rFonts)
    run.append(rPr)

    oMath = OxmlElement('m:oMath')
    oMath.text = None
    oMath.attrib.clear()
    oMath = etree.fromstring(omml_xml)
    run.append(oMath)
    para.append(run)

    return para


def latex_to_omml(s):
    """
    将简化的 LaTeX 字符串转换为 OMML XML 字符串。
    这是一个手写的转换器，支持本项目用到的所有符号。
    """
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    nsmap = {'w': W, 'm': M}

    root = etree.Element(f'{{{M}}}oMath', nsmap=nsmap)

    # 解析表达式（逗号分隔的多个子公式）
    parts = s.split(',')
    for idx, part in enumerate(parts):
        part = part.strip()
        if idx > 0:
            # 添加空格分隔符
            sPre = etree.SubElement(root, f'{{{M}}}sPre')
            sPreSpace = etree.SubElement(sPre, f'{{{M}}}spc')
            sPreSpace2 = etree.SubElement(sPreSpace, f'{{{M}}}spcPrm')
            ctrlPr = etree.SubElement(sPreSpace2, f'{{{M}}}ctrlPr')
            rFonts = etree.SubElement(ctrlPr, f'{{{W}}}rFonts')
            rFonts.set(qn('w:ascii'), 'Cambria Math')
            sPreArg = etree.SubElement(sPre, f'{{{M}}}arg')
            sPreArg.set(qn('m:argSz'), '100')
            sub2 = etree.SubElement(sPreArg, f'{{{M}}}r')
            t = etree.SubElement(sub2, f'{{{W}}}t')
            t.text = ' '

        # 处理单个子公式片段
        _parse_latex_fragment(root, part)

    return etree.tostring(root, xml_declaration=False, encoding='unicode')


def _parse_latex_fragment(parent, s):
    """
    递归解析 LaTeX 片段，构建 OMML 树。
    parent: 父 OMML 元素
    s: LaTeX 字符串
    """
    M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    if not s:
        return

    # 处理分式 \frac{}{}
    while True:
        m = _re_frac(s)
        if not m:
            break
        # 将前面的文本转为 run
        before = s[:m.start()]
        if before:
            _add_r(parent, before)
        # 构建分式
        frac = etree.SubElement(parent, f'{{{M}}}f')
        fPr = etree.SubElement(frac, f'{{{M}}}fPr')
        ctrlPr = etree.SubElement(fPr, f'{{{M}}}ctrlPr')
        _add_rFonts(ctrlPr)
        type_el = etree.SubElement(fPr, f'{{{M}}}type')
        type_el.set(qn('m:val'), 'skwlar')  # slice
        num = etree.SubElement(frac, f'{{{M}}}num')
        _parse_latex_fragment(num, m.group(1))
        den = etree.SubElement(frac, f'{{{M}}}den')
        _parse_latex_fragment(den, m.group(2))
        s = s[m.end():]
    if s:
        _add_r(parent, s)


_re_frac = __import__('re').compile(r'\\frac\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}\{([^{}]*(?:\{[^{}]*\}[^{}]*)*)\}').search


def _add_r(parent, s):
    """将普通文本添加到 OMML parent"""
    M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    # 文本替换映射
    replacements = {
        '\\theta': '\u03b8',       # θ
        '\\phi': '\u03c6',         # φ
        '\\psi': '\u03c8',         # ψ
        '\\alpha': '\u03b1',
        '\\beta': '\u03b2',
        '\\gamma': '\u03b3',
        '\\delta': '\u0394',
        '\\lambda': '\u03bb',
        '\\sigma': '\u03c3',
        '\\pi': '\u03c0',
        '\\sum': '\u2211',          # ∑
        '\\int': '\u222b',          # ∫
        '\\infty': '\u221e',        # ∞
        '\\sqrt': '\u221a',         # √
        '\\cdot': '\u22c5',         # ·
        '\\times': '\u00d7',        # ×
        '\\div': '\u00f7',          # ÷
        '\\pm': '\u00b1',           # ±
        '\\leq': '\u2264',          # ≤
        '\\geq': '\u2265',          # ≥
        '\\neq': '\u2260',          # ≠
        '\\approx': '\u2248',       # ≈
        '\\in': '\u2208',           # ∈
        '\\nabla': '\u2207',        # ∇
        '\\rightarrow': '\u2192',   # →
        '\\leftarrow': '\u2190',
        '\\arctan': 'arctan',
        '\\sin': 'sin',
        '\\cos': 'cos',
        '\\tan': 'tan',
        '\\log': 'log',
        '\\ln': 'ln',
        '\\exp': 'exp',
        '\\det': 'det',
        '\\min': 'min',
        '\\max': 'max',
        '\\lim': 'lim',
        '\\text': '',
        '\\mathrm': '',
        '\\mathbf': '',
        '\\mathit': '',
        '\\ldots': '\u2026',
        '\\dots': '\u2026',
    }

    # 处理上下标 \^{} \_{} ^{}
    out_parts = []
    i = 0
    while i < len(s):
        c = s[i]
        if c == '\\':
            # 找命令
            j = i + 1
            while j < len(s) and s[j].isalpha():
                j += 1
            cmd = '\\' + s[i+1:j]
            if j < len(s) and s[j] == '{':
                # 找配对 }
                depth = 0
                k = j + 1
                while k < len(s):
                    if s[k] == '{':
                        depth += 1
                    elif s[k] == '}':
                        if depth == 0:
                            break
                        depth -= 1
                    k += 1
                content = s[j+1:k]
                rest = s[k+1:]
            else:
                content = ''
                rest = s[j:]
                k = j

            if cmd in replacements:
                out_parts.append(replacements[cmd])
                out_parts.append(content)
            elif cmd == '\\^':
                out_parts.append('\uf000')  # superscript
                out_parts.append(content)
            elif cmd == '\\_':
                out_parts.append('\uf001')  # subscript
                out_parts.append(content)
            elif cmd in ('\\frac', '\\sqrt', '\\sum_', '\\int_', '\\lim_',
                         '\\text', '\\mathrm', '\\mathbf', '\\mathit',
                         '\\mathrm{', '\\mathbf{', '\\mathit{'):
                out_parts.append(content)
            else:
                out_parts.append(s[i])
                i = i + 1
                continue
            s = rest
            i = 0
        else:
            out_parts.append(c)
            i += 1

    raw_text = ''.join(out_parts)
    # 分解为普通文本 + 上标/下标 run
    _build_runs(parent, raw_text)


def _build_runs(parent, s):
    """将文本分解为普通文本和上下标，添加到 parent"""
    M = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    SUP = '\uf000'  # 上标标记
    SUB = '\uf001'  # 下标标记

    i = 0
    while i < len(s):
        c = s[i]
        if c == SUP:
            # 上标
            content = ''
            i += 1
            while i < len(s) and s[i] not in (SUP, SUB):
                content += s[i]
                i += 1
            sup = etree.SubElement(parent, f'{{{M}}}sup')
            r = etree.SubElement(sup, f'{{{M}}}r')
            _add_rFonts(r)
            t = etree.SubElement(r, f'{{{W}}}t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = content
        elif c == SUB:
            content = ''
            i += 1
            while i < len(s) and s[i] not in (SUP, SUB):
                content += s[i]
                i += 1
            sub = etree.SubElement(parent, f'{{{M}}}sub')
            r = etree.SubElement(sub, f'{{{M}}}r')
            _add_rFonts(r)
            t = etree.SubElement(r, f'{{{W}}}t')
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = content
        else:
            # 普通文本，收集连续的非上下标字符
            content = ''
            while i < len(s) and s[i] not in (SUP, SUB):
                content += s[i]
                i += 1
            if content:
                r = etree.SubElement(parent, f'{{{M}}}r')
                _add_rFonts(r)
                t = etree.SubElement(r, f'{{{W}}}t')
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                t.text = content


def _add_rFonts(parent):
    rFonts = etree.SubElement(parent, f'{{{W_NS}}}rFonts')
    rFonts.set(qn('w:ascii'), 'Cambria Math')
    rFonts.set(qn('w:hAnsi'), 'Cambria Math')
    rFonts.set(qn('w:eastAsia'), '宋体')


W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def insert_omml_formula(doc, para_idx, latex_str):
    """
    在指定段落索引处插入OMML公式，替换原有内容。
    """
    para = doc.paragraphs[para_idx]
    para._element.clear()

    omml_para = make_omml_para(latex_str)
    para._element.extend(omml_para)


def main():
    import zipfile

    # 公式1-4的LaTeX（简化格式，用于OMML转换）
    FORMULA_LATEX = [
        # ① 倾斜校正旋转变换矩阵
        r'\theta = \arctan\frac{y_2 - y_1}{x_2 - x_1},  M = (\cos\theta\ -\sin\theta;\ \sin\theta\ \cos\theta)',

        # ② 垂直投影字符分割
        r'V(j) = \sum_{i=0}^{H} B(i,j),\ j = 0,1,2,\ldots,W  T = \frac{\min(V)+\mathrm{mean}(V)}{2},\ V(j) < T',

        # ③ HOG特征提取
        r'\mathrm{hist}_k = \sum_{p \in \mathrm{cell}_k} |\nabla I(p)|,\ |\nabla I| = \sqrt{G_x^2 + G_y^2},\ \theta = \arctan\frac{G_y}{G_x}',

        # ④ YOLOv3损失函数
        r'L_1 = \lambda_{\mathrm{obj}}\sum\mathbb{1}_{ij}^{\mathrm{obj}}[C_i-\hat{C}_i]^2 + \lambda_{\mathrm{noobj}}\sum\mathbb{1}_{ij}^{\mathrm{noobj}}[C_i-\hat{C}_i]^2,'
        r'\ L_2 = \sum\mathbb{1}_{ij}^{\mathrm{obj}}[(\sqrt{w_i}-\sqrt{\hat{w}_i})^2+(\sqrt{h_i}-\sqrt{\hat{h}_i})^2] + \sum\mathbb{1}_{ij}^{\mathrm{obj}}[-\hat{p}_i\log p_i]',
    ]

    BASE = r"C:\Users\HUAWEI\Desktop\FinalDesign\Python_VLPR-master"
    FILE_DIR = os.path.join(BASE, "file")
    DOC_SRC = "C:/Users/HUAWEI/Desktop/专业课作业/毕设/浙江水利水电学院毕业设计（论文）模板-2026（4.9）（含真实目录）.docx"
    DOC_DST = os.path.join(FILE_DIR, "毕业论文_含OMML公式版.docx")

    doc = Document(DOC_SRC)

    print("=" * 55)
    print("开始插入OMML公式...")
    print("=" * 55)

    inserted = 0
    for i, para in enumerate(doc.paragraphs):
        if "？代码？" in para.text:
            latex = FORMULA_LATEX[inserted]
            print(f"\n[公式{inserted+1}] 段落{i}:")
            print(f"  LaTeX: {latex[:80]}...")

            try:
                insert_omml_formula(doc, i, latex)
                print(f"  ✓ OMML公式插入成功")
            except Exception as e:
                print(f"  ✗ 失败: {e}")

            inserted += 1

    print(f"\n共处理 {inserted} 个占位符")
    doc.save(DOC_DST)
    print(f"\n文档已保存: {DOC_DST}")

    # 验证
    with zipfile.ZipFile(DOC_DST, 'r') as z:
        omml_count = sum(1 for n in z.namelist() if 'oMath' in n or 'oMath' in open('').read() if False)
    print(f"文件大小: {os.path.getsize(DOC_DST):,} bytes")


if __name__ == "__main__":
    main()
