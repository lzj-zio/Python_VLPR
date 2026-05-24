# -*- coding: utf-8 -*-
"""
生成4个数学公式图片并插入到毕业论文Word文档中
公式由项目代码推导，经用户审核确认
"""
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================
# 4个公式（matplotlib mathtext 原生渲染，不用 bmatrix）
# ============================================================
FORMULAS = [
    # ① 倾斜校正旋转变换矩阵
    # 基于 img_math.py img_Transform() + 霍夫变换描述
    r"$\theta = \arctan\frac{y_2 - y_1}{x_2 - x_1},\quad "
    r"M = (\cos\theta\ -\sin\theta;\ \sin\theta\ \cos\theta)$",

    # ② 垂直投影字符分割
    # 基于 img_function.py 水平/垂直直方图 + img_math.py find_waves()
    # 阈值: T = (min(V) + mean(V)) / 2
    r"$V(j) = \sum_{i=0}^{H} B(i,j),\quad j = 0,1,2,\ldots,W\quad T = \frac{\min(V)+\mathrm{mean}(V)}{2}\quad V(j) < T$",

    # ③ HOG特征提取
    # 基于 img_recognition.py preprocess_hog()
    r"$\mathrm{hist}_k = \sum_{p \in \mathrm{cell}_k} |\nabla I(p)|\quad |\nabla I| = \sqrt{G_x^2 + G_y^2}\quad \theta = \arctan\frac{G_y}{G_x}$",

    # ④ YOLOv3损失函数（简化版，分两行）
    # 第一行：置信度损失 + 定位损失
    r"$L_1 = \lambda_{\mathrm{obj}}\sum\mathbb{1}_{ij}^{\mathrm{obj}}[C_i-\hat{C}_i]^2 + \lambda_{\mathrm{noobj}}\sum\mathbb{1}_{ij}^{\mathrm{noobj}}[C_i-\hat{C}_i]^2$"
    r"\quad $L_2 = \sum\mathbb{1}_{ij}^{\mathrm{obj}}[(\sqrt{w_i}-\sqrt{\hat{w}_i})^2+(\sqrt{h_i}-\sqrt{\hat{h}_i})^2] + \sum\mathbb{1}_{ij}^{\mathrm{obj}}[-\hat{p}_i\log p_i]$",
]


def make_formula_image(latex, out_path, dpi=200):
    """生成单个公式图片，白底黑字"""
    fig, ax = plt.subplots(figsize=(14, 2.0))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis('off')
    ax.text(0.5, 0.52, latex,
            fontsize=17, ha='center', va='center',
            transform=ax.transAxes)
    fig.tight_layout(pad=0.3)
    fig.savefig(out_path, dpi=dpi, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)


def make_yolov_formula(out_path, dpi=200):
    """YOLOv3损失函数（两行版本，更清晰）"""
    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(13, 3.2))

    # 第一行：置信度损失 + 边框损失
    line1 = (r"$L_1 = \lambda_{\mathrm{obj}}\sum\mathbb{1}_{ij}^{\mathrm{obj}}[C_i-\hat{C}_i]^2"
             r" + \lambda_{\mathrm{noobj}}\sum\mathbb{1}_{ij}^{\mathrm{noobj}}[C_i-\hat{C}_i]^2$")
    # 第二行：宽高损失 + 分类损失
    line2 = (r"$L_2 = \sum\mathbb{1}_{ij}^{\mathrm{obj}}"
             r"[(\sqrt{w_i}-\sqrt{\hat{w}_i})^2+(\sqrt{h_i}-\sqrt{\hat{h}_i})^2]"
             r" + \sum\mathbb{1}_{ij}^{\mathrm{obj}}[-\hat{p}_i\log p_i]$")

    for ax, txt in [(ax1, line1), (ax2, line2)]:
        ax.set_xlim(0, 1)
        ax.set_ylim(0, 1)
        ax.axis('off')
        ax.text(0.5, 0.55, txt, fontsize=15, ha='center', va='center',
                transform=ax.transAxes)

    fig.tight_layout(pad=0.3)
    fig.savefig(out_path, dpi=dpi, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    plt.close(fig)


def generate_all_formulas(file_dir):
    """生成全部4个公式图片"""
    os.makedirs(file_dir, exist_ok=True)
    paths = []

    # 公式1-3：标准单行
    for i in range(3):
        path = os.path.join(file_dir, f"formula_{i+1}.png")
        make_formula_image(FORMULAS[i], path)
        paths.append(path)
        print(f"  [公式{i+1}] -> {os.path.basename(path)}")

    # 公式4：YOLOv3损失函数（两行）
    path4 = os.path.join(file_dir, "formula_4.png")
    make_yolov_formula(path4)
    paths.append(path4)
    print(f"  [公式4] -> {os.path.basename(path4)}")

    return paths


def insert_formulas_to_docx(doc_src, formula_paths, doc_dst):
    """将4个公式图片插入到Word文档「？代码？」位置"""
    doc = Document(doc_src)
    inserted = 0

    for para in doc.paragraphs:
        if "？代码？" in para.text:
            para.clear()
            run = para.add_run()
            run.add_picture(formula_paths[inserted], width=Inches(5.5))
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"  [插入] 公式{inserted+1} -> 段落「？代码？」")
            inserted += 1

    doc.save(doc_dst)
    print(f"\n  文档已保存: {doc_dst}")
    return inserted


def main():
    BASE = r"C:\Users\HUAWEI\Desktop\FinalDesign\Python_VLPR-master"
    FILE_DIR = os.path.join(BASE, "file")
    DOC_SRC = "C:/Users/HUAWEI/Desktop/专业课作业/毕设/浙江水利水电学院毕业设计（论文）模板-2026（4.9）（含真实目录）.docx"
    DOC_DST = os.path.join(FILE_DIR, "毕业论文_含公式图片版.docx")

    print("=" * 55)
    print("Step 1: 生成公式图片")
    print("=" * 55)
    paths = generate_all_formulas(FILE_DIR)

    print("\n" + "=" * 55)
    print("Step 2: 插入公式到Word文档")
    print("=" * 55)
    n = insert_formulas_to_docx(DOC_SRC, paths, DOC_DST)

    print(f"\n完成！共插入 {n} 个公式图片")
    print(f"输出文件: {DOC_DST}")


if __name__ == "__main__":
    main()
